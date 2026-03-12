"""
Word 文档解析器

使用 python-docx 解析 Word 文件
"""
from typing import Dict, List
from pathlib import Path

from app.services.parsers.base_parser import BaseParser
from app.core.logger import setup_logger

logger = setup_logger(__name__)


class DOCXParser(BaseParser):
    """
    Word 文档解析器
    
    使用 python-docx 库进行解析
    """
    
    def get_supported_extensions(self) -> List[str]:
        """支持 DOCX 和 DOC 格式"""
        return ["docx", "doc"]
    
    def extract_text(self, file_path: Path) -> str:
        """
        从 Word 文档中提取纯文本
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            str: 提取的文本内容
        """
        self.validate_file(file_path)
        
        try:
            from docx import Document
            
            logger.info(f"开始解析 Word: {file_path.name}")
            
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            logger.info(f"✅ Word 解析完成: {len(paragraphs)} 个段落, {len(''.join(paragraphs))} 字符")
            
            return "\n\n".join(paragraphs)
            
        except ImportError:
            raise ImportError(
                "需要安装 python-docx: pip install python-docx\n"
                "或者在 requirements_kb.txt 中已包含，请运行: pip install -r requirements_kb.txt"
            )
        except Exception as e:
            logger.error(f"❌ Word 文档解析失败: {str(e)}", exc_info=True)
            raise RuntimeError(f"Word 文档解析失败: {str(e)}") from e
    
    def extract_with_structure(self, file_path: Path) -> Dict:
        """
        提取 Word 文档的结构化信息
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            dict: 包含文本、段落信息和元数据
        """
        self.validate_file(file_path)
        
        try:
            from docx import Document
            
            logger.info(f"开始结构化解析 Word: {file_path.name}")
            
            doc = Document(file_path)
            
            paragraphs = []
            full_text_parts = []
            
            # 提取段落
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                        "char_count": len(para.text)
                    })
                    full_text_parts.append(para.text)
            
            # 提取元数据
            metadata = {
                "file_type": "docx",
                "paragraph_count": len(paragraphs)
            }
            
            # 尝试提取文档属性
            try:
                core_properties = doc.core_properties
                metadata.update({
                    "title": core_properties.title or "",
                    "author": core_properties.author or "",
                    "subject": core_properties.subject or "",
                    "created": str(core_properties.created) if core_properties.created else "",
                    "modified": str(core_properties.modified) if core_properties.modified else ""
                })
            except Exception:
                pass
            
            logger.info(f"✅ Word 结构化解析完成: {len(paragraphs)} 个段落")
            
            return {
                "text": "\n\n".join(full_text_parts),
                "pages": [{"page": 1, "text": "\n\n".join(full_text_parts)}],
                "paragraphs": paragraphs,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"❌ Word 文档结构化解析失败: {str(e)}", exc_info=True)
            raise RuntimeError(f"Word 文档结构化解析失败: {str(e)}") from e
